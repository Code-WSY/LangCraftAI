import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
from tkinter import font
import openai
from rich import print
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re
import os

def get_result_path():
    #获取该脚本所在的目录
    current_path = os.path.dirname(os.path.abspath(__file__))
    #打开上级目录
    current_path = os.path.dirname(current_path)
    #打开上级目录的result目录
    result_path= os.path.join(current_path, 'result')
    #没有result目录就创建
    if not os.path.exists(result_path):
        os.makedirs(result_path)
    return result_path

def cd_result_path():
    cd_result_path = get_result_path()
    if os.name == 'nt':
        os.system(f"start {cd_result_path}")
    else:
        os.system(f"open {cd_result_path}")
    return cd_result_path



def READ_ME():

    # 定义系统字体
    System_font = font.Font(family='楷体', size=14)
    # 创建控制台对象
    text_dict = {
        "                   --------❤欢迎使用善姬的AI辅助工具❤-------- \n": "red",
        "                                   v-1.0.0\n": "red",
        "使用说明：\n\n": "white",
        "    1. 选择文件：需要翻译/润色的文件(目前支持文本格式文件和word文档)\n": "white",
        "    2. 选择提示词文件：即翻译/润色要求（目前支持文本格式文件和word文档）\n": "white",
        "    3. 选择是否需要进行段段对照排版\n": "white",
        "    4. 选择分割符号：默认为分割线\n": "white",
        "       注意：希望按照段落进行翻译/润色，可以选择换行符进行分割\n": "white",
        "             如果每段之间有一个空行，应该选择两个换行符\n": "white",
        "             如果每段之间没有空行，应该选择一个换行符\n": "white",
        "    5. 选择翻译模型：建议选择默认的模型\n": "white",
        "    6. 提交进行翻译\n\n": "white",
        "注意事项:\n\n": "orange",
        "    - 每翻译完一个部分（section），翻译/润色内容会自动添加到文本文件中，\n": "yellow",
        "      文件名为: 原文件名+_translated.txt。\n\n": "yellow",
        "    - 若翻译过程中出现中断\n": "yellow",
        "      可以将原文件中已经翻译/润色完成的内容删除，重新运行程序即可。\n\n": "yellow",
        "    - 全部翻译/润色完成后会自动将翻译内容保存为word文档，\n": "yellow",
        "      文件名为: 原文件名+_translated.docx。\n\n": "yellow",
       f"    - 保存目录：{get_result_path()}\n": "yellow",
    }
    #console.print(text)
    return text_dict


def read_API_key(filename='API_KEY'):
    # 读取 API 密钥,第一行是key,第二行是base_url(如果有且不为空)
    with open(filename, 'r', encoding='utf-8') as file:
        api_key = file.readline().strip()
        base_url = file.readline().strip()
        if base_url:
            openai.base_url = base_url
        #print(f"API_KEY: {api_key}")
        #print(f"base_url: {openai.base_url}")

    return api_key


def user_messages(text):
    role = "user"
    message = text
    usermsg = {"role": role, "content": message}
    return usermsg


def system_messages(text):
    role = "system"
    message = text
    sysmsg = {"role": role, "content": message}
    return sysmsg


def read_and_split(filename, split_symbol):
    #如果是word文档
    if filename.endswith('.docx') or filename.endswith('.doc'):
        #读取word文档至文本
        doc = Document(filename)
        content = ""
        for para in doc.paragraphs:  #遍历段落，这里的段落是指word文档中的段落
            content += para.text + '\n\n'
        # 分割文本
        sections = content.split(split_symbol)
        #删除sections中的空字符串
        sections = [section for section in sections if section]
        return sections
    # 打开文件并读取内容
    with open(filename, 'r', encoding='utf-8') as file:
        content = file.read()
    # 使用连续两个换行符分割文本
    sections = content.split(split_symbol)
    # 返回分割后的列表
    return sections


def select_file(file_path):
    filename = os.path.basename(file_path)
    save_name = f'{filename.split(".")[0]}_translated.txt'
    docx_name = f'{filename.split(".")[0]}_translated.docx'
    return filename, save_name, docx_name


def read_text(file_path):
    if file_path.endswith('.docx') or file_path.endswith('.doc'):
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + '\n'
        return text
    # 选择后读取文件内容
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()  # 返回的是字符串
    return text


def save_text_to_file(text, filename):
    filename = os.path.join(get_result_path(), filename)
    #补充写入
    with open(filename, 'a', encoding='utf-8') as file:
        file.write(text)
    return None


def save_to_docx(text_name, save_name):
    save_name = os.path.join(get_result_path(), save_name)
    text_name = os.path.join(get_result_path(), text_name)
    # 创建 Word 文档
    doc = Document()
    # 设置整个文档的默认英文字体
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    # 打开 Markdown 文件并读取内容
    with open(text_name, 'r', encoding='utf-8') as file:
        content = file.read()
    # 将内容按段落分开添加到文档中
    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        p = doc.add_paragraph()
        # 通过正则表达式检测段落是否包含中文字符
        if re.search('[\u4e00-\u9fa5]', paragraph):
            # 为段落设置中文字体样式
            run = p.add_run(paragraph)
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        else:
            # 为段落设置英文字体样式
            run = p.add_run(paragraph)
            run.font.name = 'Times New Roman'
    # 保存 Word 文档
    doc.save(save_name)
    # 输出保存成功的信息
    print("已保存为word文档：" + save_name)


def get_gpt_response(output_text, chat_history, user_message, model, temperature, max_tokens=4096):
    chat_history.append(user_message)
    response = openai.chat.completions.create(
        model=model,
        messages=chat_history,
        temperature=temperature,
        max_tokens=max_tokens,
        stream=True
    )
    message = " "
    text = model + ": \n"
    output_text.config(state=tk.NORMAL)
    output_text.see(tk.END)
    output_text.insert(tk.END, text, 'English')
    output_text.config(state=tk.DISABLED)
    output_text.update()
    for idx, chunk in enumerate(response):
        if len(chunk.choices) == 0:
            continue
        else:
            chunk_message = chunk.choices[0].delta
        if not chunk_message.content:
            continue
        message += chunk_message.content
        #实时加入文本框
        output_text.config(state=tk.NORMAL)
        output_text.insert(tk.END, chunk_message.content, 'Chinese')
        #移动到最后
        output_text.see(tk.END)
        output_text.config(state=tk.DISABLED)
        output_text.update()
    chat_history.append({'role': 'assistant', 'content': message})
    return chat_history


def chat_gpt(output_text, filename, save_name, split_symbol, chat_history, is_cross_reference, model, temperature,
             max_tokens,remember_length):
    sections = read_and_split(filename, split_symbol=split_symbol)
    idx = 0
    # 清屏
    output_text.config(state=tk.NORMAL)
    output_text.delete(1.0, tk.END)
    output_text.config(state=tk.DISABLED)
    for section in sections:
        idx += 1
        print(f"Section {idx}/{len(sections)}:")
        print(section)
        output_text.config(state=tk.NORMAL)
        output_text.insert(tk.END, f"\nSection {idx}/{len(sections)}:\n", 'English')
        output_text.config(state=tk.DISABLED)
        #print(section)
        if len(chat_history) >= 1 and len(chat_history) <= 2 * remember_length + 1:  #1+2*2=5
            chat_history = get_gpt_response(output_text, chat_history, user_messages(section),
                                            model, temperature, max_tokens)
            if is_cross_reference:
                save_text_to_file(section + '\n', save_name)
                save_text_to_file(chat_history[-1]['content'], save_name)
                save_text_to_file('\n' + split_symbol + '\n', save_name)
            else:
                # 不保存原文
                save_text_to_file(chat_history[-1]['content'], save_name)
                save_text_to_file('\n' + split_symbol + '\n', save_name)
        else:
            # 只取第一个和后面的几个
            chat_history = chat_history[:1] + chat_history[-2 * remember_length:]
            chat_history = get_gpt_response(output_text, chat_history, user_messages(section),
                                            model, temperature, max_tokens)
            save_text_to_file(section + '\n', save_name)
            save_text_to_file(chat_history[-1]['content'], save_name)
            save_text_to_file('\n' + split_symbol + '\n', save_name)
        print('-' * 40)
        output_text.config(state=tk.NORMAL)
        #换行
        output_text.insert(tk.END, '\n', 'Chinese')
        output_text.insert(tk.END, '-' * 40, 'Chinese')
        output_text.insert(tk.END, '\n', 'Chinese')
        output_text.config(state=tk.DISABLED)


def output_display(app, text=None):
    # 设置字体,颜色：
    Chinese_font = font.Font(family='楷体', size=13)
    English_font = font.Font(family='Times New Roman', size=12, weight='bold')
    System_font = font.Font(family='楷体', size=14)
    # 创建一个文本显示框，不可编辑，深灰黑：背景色，白色：前景色，加粗
    #好看的颜色：dark slate gray
    output_text = tk.Text(app, height=30, width=200,  # 如果按照百分比设置大小，会随着窗口大小变化，用参数：height=0.5, width=0.5
                          bg='#003366',
                          fg='#F0F8FF',  #003366:深灰黑, F0F8FF:白色
                          wrap=tk.WORD, font=Chinese_font)
    # 使文本框可滚动
    output_scroll = tk.Scrollbar(app, command=output_text.yview)
    output_scroll.pack(side=tk.RIGHT, fill=tk.Y)

    # 配置标签使用不同的字体
    output_text.tag_configure('Chinese', font=Chinese_font, foreground='#F0F8FF',
                              background='#003366', selectbackground='green')
    output_text.tag_configure('English', font=English_font, foreground='red',
                              background='#003366', selectbackground='green')
    #output_text.tag_configure('system', font=System_font,
    #                          foreground='green', background='#003366',selectbackground='green')

    # 配置标签及其样式
    for style in set(text.values()):  # 使用set去除重复的样式
        output_text.tag_configure(style, foreground=style, font=System_font)

    # 插入文本并应用相应的标签
    for text, style in text.items():
        output_text.insert(tk.END, text, style)
        output_text.config(state=tk.NORMAL)

    # 设置文本框为不可编辑
    output_text.config(state=tk.DISABLED)
    return output_text


def submit(output_text, file_path, promotion_path, is_cross_reference, split_symbol, model,remember_length):
    #检查文件是否选择
    if file_path == "选择文件":
        output_text.config(state=tk.NORMAL)
        #清空界面
        output_text.delete(1.0, tk.END)
        output_text.insert(tk.END, "未选择文件！\n")
        return None
    if promotion_path == "选择提示词文件":
        output_text.config(state=tk.NORMAL)
        output_text.delete(1.0, tk.END)
        output_text.insert(tk.END, "未选择提示词文件！\n")
        return None
    #检查结束，开始翻译
    #选择提示内容
    promotion = read_text(promotion_path)
    #选择分割符号
    split_symbol = select_split_symbol_value(split_symbol)
    #选择是否段段对照
    is_cross_reference = is_cross_reference
    #选择模型
    model = model
    _, save_name, docx_name = select_file(file_path)
    #开始翻译
    chat_history = [system_messages(promotion)]
    chat_gpt(output_text, file_path, save_name, split_symbol, chat_history, is_cross_reference, model,
             temperature=0.5, max_tokens=4096, remember_length=remember_length)
    save_to_docx(save_name, docx_name)
    return None


def GUI(choices_model, remember_length):
    app = tk.Tk()
    app.title("善姬的翻译器")
    app.geometry('800x800')
    app.resizable(width=True, height=False)
    file_path = tk.StringVar()
    file_path.set("选择文件")
    promotion_path = tk.StringVar()
    promotion_path.set("选择提示词文件")
    split_symbol = tk.StringVar()

    is_cross_reference = tk.IntVar()
    select_model = tk.StringVar()
    #---------------------------------------------------------------------
    #文件菜单
    menu_bar = tk.Menu(app)
    file_menu = tk.Menu(menu_bar, tearoff=0)
    menu_bar.add_cascade(label="文件", menu=file_menu)
    file_menu.add_command(label="目录", command=cd_result_path)
    file_menu.add_separator()  #分割线
    file_menu.add_command(label="退出", command=app.quit)
    app.config(menu=menu_bar)
    #--------------------------------------------------------------
    # 输出文本框
    # 设置字体
    chinese_font = font.Font(family='仿宋', size=12)  # 中文仿宋字体
    english_font = font.Font(family='Times New Roman', size=12)  # 英文Times New Roman字体
    init_text = READ_ME()
    output_text = output_display(app, text=init_text)
    output_text.pack()

    #--------------------------------------------------------------
    # 文件选择按钮，居中：tk.TOP, tk.BOTTOM, tk.LEFT, tk.RIGHT
    def select_file():
        file_paths = filedialog.askopenfilename()
        if file_paths:
            # 这里可以添加你需要对文件执行的操作
            print(f"选择的文件: {file_paths}")
            # label
            file_label.config(text="原文件:" + file_paths.split('/')[-1])
            file_path.set(file_paths)

    def select_promotion():
        promo_path = filedialog.askopenfilename()
        if promo_path:
            # 这里可以添加你需要对文件执行的操作
            print(f"选择的文件: {promo_path}")
            # label
            promo_label.config(text="提示文件:" + promo_path.split('/')[-1])
            promotion_path.set(promo_path)

    file_button = tk.Button(app,
                            command=select_file, width=30, height=1,
                            textvariable=file_path)
    file_button.pack(side=tk.TOP, padx=1, pady=1)  #PADX: 水平间距，PADY: 垂直间距

    file_label = tk.Label(app, text="")
    file_label.pack(pady=1)

    promo_button = tk.Button(app,
                             command=select_promotion, width=30, height=1,
                             textvariable=promotion_path)

    promo_button.pack(side=tk.TOP, padx=1, pady=1)
    promo_label = tk.Label(app, text="")
    promo_label.pack(pady=1)
    #--------------------------------------------------------------
    #选择按钮：是否段段对照
    cross_reference_button = tk.Checkbutton(app, text="段段对照排版", variable=is_cross_reference)
    cross_reference_button.pack(padx=10, pady=10)
    #--------------------------------------------------------------
    #分割符号
    split_symbol_dict = choice_split_symbol_key()
    split_symbol = ttk.Combobox(app, values=list(split_symbol_dict.keys()),
                                state='readonly',
                                textvariable=split_symbol)
    split_symbol.pack(padx=10, pady=10)
    split_symbol.current(2)  # 设置默认值
    #--------------------------------------------------------------
    # 选择模型列表
    model = ttk.Combobox(app, values=choices_model, state='readonly', textvariable=select_model)
    model.pack(padx=10, pady=10)
    model.current(16)  # 设置默认值
    submit_button = tk.Button(app, text="提交",
                              command=lambda: submit(output_text, file_path.get(), promotion_path.get(),
                                                     is_cross_reference.get(), split_symbol.get(), select_model.get()
                                                        ,remember_length,
                                                     )
                              , width=10, height=1)

    split_symbol.pack(padx=10, pady=10)
    submit_button.pack(side=tk.TOP, padx=10, pady=10, anchor=tk.CENTER)  # anchor: 锚定位置
    return app


def choice_model_list(model_list):
    model_list = [model.id for model in model_list]
    #filter_model=['babbage-002','code-davinci-edit-001',
    #              'dall-e','gpt','davinci-002','text','tts','whisper']
    filter_model = ['gpt']
    #如果是filter_model中的开头的模型，就显示
    user_model_list = []
    for model in model_list:
        for filter in filter_model:
            if model.startswith(filter):
                user_model_list.append(model)
    return user_model_list


def choice_split_symbol_key():
    split_symbol_dict = {
        '两个换行符': '\\n\\n',
        '一个换行符': '\\n',
        '分割线': '---',
        '句号': '。',
        '分号': '；',
        '逗号': '，',
        '问号': '？',
        '感叹号': '！',
        '冒号': '：'
    }
    return split_symbol_dict


def select_split_symbol_value(split_symbol):
    split_symbol_dict = choice_split_symbol_key()
    value = split_symbol_dict[split_symbol]
    print(f"分割符号：{value}")
    # 对获取到的value进行转义
    value = value.encode().decode('unicode_escape')
    return value


if __name__ == '__main__':
    remember_length = 2
    openai.api_key = read_API_key()
    choices_model = choice_model_list(openai.models.list())
    app = GUI(choices_model, remember_length)
    app.mainloop()
