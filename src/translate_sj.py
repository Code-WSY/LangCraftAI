import openai
from rich.console import Console
from rich.text import Text
from rich import print
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog
import os


def read_API_key(filename='API_KEY'):
    # 读取 API 密钥,只读取第一行
    with open(filename, 'r', encoding='utf-8') as file:
        api_key = file.readline().strip()
    return api_key
def read_and_split(filename, split_symbol):
    # 打开文件并读取内容
    with open(filename, 'r', encoding='utf-8') as file:
        content = file.read()
    # 使用连续两个换行符分割文本
    sections = content.split(split_symbol)
    # 返回分割后的列表
    return sections

def save_to_docx(text_name, save_name):
    # 创建 Word 文档
    doc = Document()
    # 设置整个文档的默认英文字体
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    # 打开 Markdown 文件并读取内容
    with open(text_name, 'r', encoding='utf-8') as file:
        content = file.read()
    # 将内容按段落分开添加到文档中
    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        p = doc.add_paragraph()
        # 通过正则表达式检测段落是否包含中文字符
        if re.search('[\u4e00-\u9fa5]', paragraph):
            # 为段落设置中文字体样式
            run = p.add_run(paragraph)
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        else:
            # 为段落设置英文字体样式
            run = p.add_run(paragraph)
            run.font.name = 'Times New Roman'
    # 保存 Word 文档
    doc.save(save_name)
    # 输出保存成功的信息
    print("已保存为word文档：" + save_name)

def select_file():
    print("请选择要翻译的文件：")
    # 创建一个 Tkinter 根窗口并隐藏
    root = tk.Tk()
    root.withdraw()
    #获取当前文件夹路径
    file_path = os.getcwd()
    # 显示文件选择对话框并获取选择的文件路径
    file_path = filedialog.askopenfilename(initialdir=file_path)
    if file_path:
        filename = os.path.basename(file_path)
        save_name = f'{filename.split(".")[0]}_translated.{filename.split(".")[1]}'
        docx_name = f'{filename.split(".")[0]}_translated.docx'
    else:
        raise ValueError("No file selected.")

    return filename, save_name, docx_name
def promotion_file():
    print("请选择提示内容文件：")
    # 创建一个 Tkinter 根窗口并隐藏
    root = tk.Tk()
    root.withdraw()
    file_path = os.getcwd()
    # 显示文件选择对话框并获取选择的文件路径
    file_path = filedialog.askopenfilename(initialdir=file_path)
    if not file_path:
        raise ValueError("No file selected.")
    # 选择后读取文件内容
    with open(file_path, 'r', encoding='utf-8') as file:
        promotion = file.read() # 返回的是字符串
    print("已选择提示内容：" + promotion)
    return promotion

def cross_reference():
    #选择是否段段对照
    is_cross_reference = input("是否进行段段对照排版？\n1.是\n2.否\n>>> ")
    #检查格式
    while True:
        if is_cross_reference == "1":
            is_cross_reference = True
            break
        elif is_cross_reference == "2":
            is_cross_reference = False
            break
        else:
            is_cross_reference = input("请输入正确的选项：\n1.是\n2.否\n>>> ")

    return is_cross_reference


def select_split_symbol():
    tips_doc = ("提示：\n"
                "如果您的文本每段中间有空一行，请输入: \\n\\n\n"
                "如果您的文本每段直接无空行，请输入: \\n\n")
    split_symbol = input(f"\n{tips_doc}\n请输入分割符号：>>> ")
    # 分割符号不能是空
    while split_symbol == "":
        split_symbol = input("分割符号不能为空，请重新输入：>>> ")
    # 分隔符号需要转义
    split_symbol = split_symbol.encode().decode('unicode_escape')
    print(f"已选择分割符号：{split_symbol}")
    return split_symbol



def user_messages(text):
    role = "user"
    message = text
    usermsg = {"role": role, "content": message}
    return usermsg
def system_messages(text):
    role = "system"
    message = text
    sysmsg = {"role": role, "content": message}
    return sysmsg

def save_text_to_file(text, filename):
    #补充写入
    with open(filename, 'a', encoding='utf-8') as file:
        file.write(text)

    return None

class GPT:
    def __init__(self, api_key, system_message="",is_cross_reference=False):
        self.api_key = api_key
        self.base_url = "https://oneapi.xty.app/v1/"
        self.api_key = self.api_key
        self.openai = openai.OpenAI(api_key=self.api_key, base_url=self.base_url)
        self.model = "gpt-4-0125-preview"
        self.temperature = 0.4
        self.max_tokens = 4096
        self.system_message = system_messages(system_message)
        self.chat_history = [self.system_message]
        self.response = ""
        self.user_message = ""
        self.model_list = self.openai.models.list()
        self.is_cross_reference = is_cross_reference
    def choice_model(self):
        model_list = [model.id for model in self.model_list]
        #filter_model=['babbage-002','code-davinci-edit-001',
        #              'dall-e','gpt','davinci-002','text','tts','whisper']
        filter_model = ['gpt']
        #如果是filter_model中的开头的模型，就显示
        user_model_list = []
        for model in model_list:
            for filter in filter_model:
                if model.startswith(filter):
                    user_model_list.append(model)
        print("模型列表：")
        for idx, model in enumerate(user_model_list):
            print(f"{idx + 1}. {model}")
        print("建议：选择17")
        model_id = input("请选择模型>>> ")
        self.model = user_model_list[int(model_id) - 1]
        print(f"已选择模型：{self.model}")
        return self.model
    def choice_model_list(self):
        model_list = [model.id for model in self.model_list]
        #filter_model=['babbage-002','code-davinci-edit-001',
        #              'dall-e','gpt','davinci-002','text','tts','whisper']
        filter_model = ['gpt']
        #如果是filter_model中的开头的模型，就显示
        user_model_list = []
        for model in model_list:
            for filter in filter_model:
                if model.startswith(filter):
                    user_model_list.append(model)
        return user_model_list

    def get_gpt_response(self, chat_history, user_message):
        console = Console()
        chat_history.append(user_message)
        self.response = self.openai.chat.completions.create(
            model=self.model,
            messages=chat_history,
            temperature=self.temperature,
            max_tokens=self.max_tokens,
            stream=True
        )
        message = " "
        console.print(Text(self.model + ": ", style="bold blue"))
        for idx, chunk in enumerate(self.response):
            if len(chunk.choices) == 0:
                continue
            else:
                chunk_message = chunk.choices[0].delta
            if not chunk_message.content:
                continue
            message += chunk_message.content
            text = Text(chunk_message.content, style="bold green")
            console.print(text, end="")
        chat_history.append({'role': 'assistant', 'content': message})
        console.print()
        return chat_history

    def chat_gpt(self, filename):
        sections = read_and_split(filename, split_symbol=split_symbol)
        idx = 0
        self.choice_model()
        #清屏
        console = Console()
        console.clear()
        for section in sections:
            idx += 1
            print(f"Section {idx}/{len(sections)}:")
            print(section)
            if len(self.chat_history) >= 1 and len(self.chat_history) <=3:
                self.chat_history=self.get_gpt_response(self.chat_history, user_messages(section))
                if self.is_cross_reference:
                    save_text_to_file(section+'\n', save_name)
                    save_text_to_file(self.chat_history[-1]['content'], save_name)
                    save_text_to_file('\n'+split_symbol+'\n', save_name)
                else:
                    # 不保存原文
                    save_text_to_file(self.chat_history[-1]['content'], save_name)
                    save_text_to_file('\n'+split_symbol+'\n', save_name)
            else:
                #只取第一个和后面的4个
                self.chat_history = self.chat_history[:1] + self.chat_history[-2:]
                self.chat_history = self.get_gpt_response(self.chat_history, user_messages(section))
                save_text_to_file(section+'\n', save_name)
                save_text_to_file(self.chat_history[-1]['content'], save_name)
                save_text_to_file('\n'+split_symbol+'\n', save_name)
            print('-' * 40)
            #休息一秒
            import time
            time.sleep(3)

if __name__ == '__main__':
    openai.api_key = read_API_key()
    # 选择提示内容
    promotion = promotion_file()
    # 选择翻译文件
    filename, save_name, docx_name = select_file()
    # 选择分割符号
    split_symbol = select_split_symbol()
    #选择是否段段对照
    is_cross_reference = cross_reference()
    gpt = GPT(api_key=openai.api_key, system_message=promotion,is_cross_reference=is_cross_reference)
    gpt.chat_gpt(filename)
    save_to_docx(save_name, docx_name)






